# Feature: Multi-format Document Processing
# Description: Process various document formats (PDF, DOCX, TXT, etc.) for text extraction
# Library: Django, python-docx, PyPDF2, python-magic

# 1. models.py - Document storage model
from django.db import models
import uuid

class ProcessedDocument(models.Model):
    id = models.UUIDField(primary_key=True, default=uuid.uuid4)
    title = models.CharField(max_length=200)
    filename = models.CharField(max_length=200)
    file_type = models.CharField(max_length=10)
    content = models.TextField()
    word_count = models.IntegerField(default=0)
    token_estimate = models.IntegerField(default=0)
    created_at = models.DateTimeField(auto_now_add=True)
    
    class Meta:
        ordering = ['-created_at']

# 2. document_processors.py - Format-specific processors
import os
import magic
from docx import Document as DocxDocument
import PyPDF2
from io import BytesIO

class DocumentProcessor:
    
    SUPPORTED_FORMATS = {
        'txt': ['text/plain'],
        'pdf': ['application/pdf'],
        'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'html': ['text/html'],
        'md': ['text/plain', 'text/x-markdown'],
    }
    
    @staticmethod
    def detect_file_type(file_path):
        """Detect file type using python-magic"""
        mime_type = magic.from_file(file_path, mime=True)
        
        for ext, mime_types in DocumentProcessor.SUPPORTED_FORMATS.items():
            if mime_type in mime_types:
                return ext
        
        # Fallback to extension
        return os.path.splitext(file_path)[1][1:].lower()
    
    @staticmethod
    def process_txt(file_path):
        """Process text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return {
                'success': True,
                'content': content,
                'word_count': len(content.split()),
                'error': None
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def process_pdf(file_path):
        """Process PDF files"""
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content.append(page.extract_text())
            
            full_content = '\n'.join(content)
            return {
                'success': True,
                'content': full_content,
                'word_count': len(full_content.split()),
                'error': None
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    @staticmethod
    def process_docx(file_path):
        """Process DOCX files"""
        try:
            doc = DocxDocument(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            full_content = '\n'.join(content)
            return {
                'success': True,
                'content': full_content,
                'word_count': len(full_content.split()),
                'error': None
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    @classmethod
    def process_file(cls, file_path, filename):
        """Main processing method that routes to specific processors"""
        file_type = cls.detect_file_type(file_path)
        
        processors = {
            'txt': cls.process_txt,
            'md': cls.process_txt,
            'html': cls.process_txt,
            'pdf': cls.process_pdf,
            'docx': cls.process_docx,
        }
        
        if file_type not in processors:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_type}'
            }
        
        result = processors[file_type](file_path)
        if result['success']:
            # Estimate token count (rough approximation: 1 token ≈ 4 characters)
            result['token_estimate'] = len(result['content']) // 4
        
        return result

# 3. views.py - Document upload and processing endpoint
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
from django.utils.decorators import method_decorator
from django.core.files.storage import default_storage
import tempfile
import os

@method_decorator(csrf_exempt, name='dispatch')
class DocumentUploadView(View):
    def post(self, request):
        try:
            if 'file' not in request.FILES:
                return JsonResponse({'error': 'No file provided'}, status=400)
            
            uploaded_file = request.FILES['file']
            filename = uploaded_file.name
            
            # Save file temporarily
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_path = temp_file.name
            
            try:
                # Process the document
                result = DocumentProcessor.process_file(temp_path, filename)
                
                if not result['success']:
                    return JsonResponse({'error': result['error']}, status=400)
                
                # Save processed document to database
                doc = ProcessedDocument.objects.create(
                    title=os.path.splitext(filename)[0],
                    filename=filename,
                    file_type=DocumentProcessor.detect_file_type(temp_path),
                    content=result['content'],
                    word_count=result['word_count'],
                    token_estimate=result['token_estimate']
                )
                
                return JsonResponse({
                    'success': True,
                    'document': {
                        'id': str(doc.id),
                        'title': doc.title,
                        'filename': doc.filename,
                        'file_type': doc.file_type,
                        'word_count': doc.word_count,
                        'token_estimate': doc.token_estimate,
                        'content_preview': doc.content[:200] + '...' if len(doc.content) > 200 else doc.content
                    }
                })
                
            finally:
                # Clean up temporary file
                os.unlink(temp_path)
                
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

class DocumentListView(View):
    def get(self, request):
        try:
            documents = ProcessedDocument.objects.all()[:20]  # Last 20 documents
            
            doc_list = []
            for doc in documents:
                doc_list.append({
                    'id': str(doc.id),
                    'title': doc.title,
                    'filename': doc.filename,
                    'file_type': doc.file_type,
                    'word_count': doc.word_count,
                    'token_estimate': doc.token_estimate,
                    'created_at': doc.created_at.isoformat()
                })
            
            return JsonResponse({'documents': doc_list})
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

# 4. urls.py
from django.urls import path
from .views import DocumentUploadView, DocumentListView

urlpatterns = [
    path('documents/upload/', DocumentUploadView.as_view()),
    path('documents/', DocumentListView.as_view()),
]